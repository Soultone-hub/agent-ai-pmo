"""
Service d'export des analyses générées (PDF / Word / Excel).

Chaque type d'analyse (`document`, `risks`, `kpi`, `copil`) est rendu dans un
document mis en forme et directement exploitable par un PMO :
- PDF  : généré via l'API `Story` de PyMuPDF (HTML/CSS → PDF), sans dépendance externe.
- DOCX : généré via python-docx.
- XLSX : généré via openpyxl.

Point d'entrée : `build_export(analysis_type, result, project_name, created_at, fmt)`.
"""
from __future__ import annotations

import html
import io
from datetime import datetime
from typing import Any

import fitz  # PyMuPDF

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


# ── Métadonnées de présentation ──────────────────────────────────────────────
TYPE_LABELS: dict[str, str] = {
    "document": "Analyse documentaire",
    "risks": "Registre des risques",
    "kpi": "Tableau de bord KPI",
    "copil": "Synthèse COPIL",
}

ALLOWED_FORMATS = ("pdf", "docx", "xlsx")

MEDIA_TYPES: dict[str, str] = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

GENERATED_BY = "Généré par Agent IA PMO — Groq LLaMA 3.3-70B"

# Couleurs (hex sans #) par niveau de risque
RISK_COLORS = {
    "critique": "ef4444",
    "eleve": "f59e0b",
    "modere": "6366f1",
    "faible": "10b981",
}
RISK_LABELS = {
    "critique": "Critique",
    "eleve": "Élevé",
    "modere": "Modéré",
    "faible": "Faible",
}

# Couleurs par statut KPI
KPI_COLORS = {
    "vert": "10b981",
    "orange": "f59e0b",
    "rouge": "ef4444",
    "inconnu": "9ca3af",
}
KPI_LABELS = {
    "vert": "Vert", "orange": "Orange", "rouge": "Rouge", "inconnu": "Inconnu",
}

ALERTE_LABELS = {
    "retard": "Retard", "budget": "Budget", "ressource": "Ressource", "risque": "Risque",
}
PRIORITE_LABELS = {
    "critique": "Critique", "haute": "Haute", "moyenne": "Moyenne", "faible": "Faible",
}
STATUT_ACTION_LABELS = {
    "a_faire": "À faire", "en_cours": "En cours", "termine": "Terminé", "bloque": "Bloqué",
}

BRAND = "4f46e5"  # indigo


# ── Utilitaires ───────────────────────────────────────────────────────────────
def _fmt_date(created_at: Any) -> str:
    if isinstance(created_at, datetime):
        return created_at.strftime("%d/%m/%Y à %H:%M")
    if created_at:
        try:
            return datetime.fromisoformat(str(created_at)).strftime("%d/%m/%Y à %H:%M")
        except ValueError:
            return str(created_at)
    return datetime.now().strftime("%d/%m/%Y à %H:%M")


def _safe_filename(project_name: str, analysis_type: str, fmt: str) -> str:
    base = f"{TYPE_LABELS.get(analysis_type, analysis_type)}_{project_name}"
    keep = []
    for ch in base:
        if ch.isalnum() or ch in (" ", "-", "_"):
            keep.append(ch)
        elif ch in "àâäéèêëïîôöùûüç":
            keep.append(ch)
    cleaned = "".join(keep).strip().replace(" ", "_")
    return f"{cleaned or 'export'}.{fmt}"


def _as_list(value: Any) -> list:
    return value if isinstance(value, list) else []


# ══════════════════════════════════════════════════════════════════════════════
#  GÉNÉRATION PDF (PyMuPDF Story / HTML)
# ══════════════════════════════════════════════════════════════════════════════
def _esc(text: Any) -> str:
    return html.escape(str(text if text is not None else ""))


def _html_header(label: str, project_name: str, date_str: str) -> str:
    return f"""
    <div style="border-bottom:2px solid #{BRAND}; padding-bottom:6px; margin-bottom:14px;">
      <p style="color:#{BRAND}; font-size:9pt; letter-spacing:1px; margin:0;">{_esc(label.upper())}</p>
      <h1 style="font-size:20pt; margin:2px 0 0 0; color:#111827;">{_esc(project_name)}</h1>
      <p style="font-size:9pt; color:#6b7280; margin:4px 0 0 0;">Document généré le {_esc(date_str)}</p>
    </div>
    """


def _html_section_title(title: str) -> str:
    return (f'<p style="font-size:11pt; color:#{BRAND}; font-weight:bold; '
            f'margin:16px 0 4px 0;">{_esc(title)}</p>')


def _html_bullets(items: list) -> str:
    if not items:
        return '<p style="font-size:10pt; color:#9ca3af; font-style:italic;">Aucun élément.</p>'
    lis = "".join(f'<li style="margin-bottom:3px;">{_esc(i)}</li>' for i in items)
    return f'<ul style="font-size:10pt; color:#374151; margin:0 0 0 14px;">{lis}</ul>'


def _build_html(analysis_type: str, result: dict, project_name: str, date_str: str) -> str:
    label = TYPE_LABELS.get(analysis_type, analysis_type)
    body = [_html_header(label, project_name, date_str)]

    if analysis_type == "document":
        body.append(_html_section_title("Résumé"))
        body.append(f'<p style="font-size:10pt; color:#374151;">{_esc(result.get("resume") or "—")}</p>')
        for key, title in [
            ("objectifs", "Objectifs"),
            ("contraintes", "Contraintes"),
            ("hypotheses", "Hypothèses"),
            ("parties_prenantes", "Parties prenantes"),
            ("points_cles", "Points clés"),
        ]:
            body.append(_html_section_title(title))
            body.append(_html_bullets(_as_list(result.get(key))))

    elif analysis_type == "risks":
        risks = _as_list(result.get("risks"))
        body.append(_html_section_title(f"Synthèse — {len(risks)} risque(s) identifié(s)"))
        if result.get("resume"):
            body.append(f'<p style="font-size:10pt; color:#374151;">{_esc(result.get("resume"))}</p>')
        rows = ['<tr style="background-color:#f3f4f6;">'
                '<th align="left">ID</th><th align="left">Description</th>'
                '<th align="left">Niveau</th><th align="left">P</th>'
                '<th align="left">I</th><th align="left">Score</th>'
                '<th align="left">Catégorie</th><th align="left">Mitigation</th></tr>']
        for r in risks:
            niveau = (r or {}).get("niveau", "faible")
            color = RISK_COLORS.get(niveau, "9ca3af")
            rows.append(
                "<tr>"
                f'<td>{_esc(r.get("id"))}</td>'
                f'<td>{_esc(r.get("description"))}</td>'
                f'<td><b style="color:#{color};">{_esc(RISK_LABELS.get(niveau, niveau))}</b></td>'
                f'<td>{_esc(r.get("probabilite"))}</td>'
                f'<td>{_esc(r.get("impact"))}</td>'
                f'<td>{_esc(r.get("score"))}</td>'
                f'<td>{_esc(r.get("categorie"))}</td>'
                f'<td>{_esc(r.get("mitigation"))}</td>'
                "</tr>"
            )
        body.append(
            '<table border="1" cellpadding="4" cellspacing="0" '
            'style="font-size:8.5pt; color:#374151; width:100%; border-color:#d1d5db;">'
            + "".join(rows) + "</table>"
        )

    elif analysis_type == "kpi":
        kpis = _as_list(result.get("kpis"))
        score = (result.get("score_global") or {})
        body.append(_html_section_title("Score global de santé"))
        body.append(
            f'<p style="font-size:14pt; color:#{BRAND}; font-weight:bold; margin:0;">'
            f'{_esc(score.get("valeur", 0))}/100</p>'
            f'<p style="font-size:10pt; color:#374151;">{_esc(score.get("interpretation") or "—")}</p>'
        )
        if result.get("resume"):
            body.append(f'<p style="font-size:10pt; color:#6b7280;">{_esc(result.get("resume"))}</p>')
        body.append(_html_section_title(f"Indicateurs — {len(kpis)} KPI"))
        rows = ['<tr style="background-color:#f3f4f6;">'
                '<th align="left">Indicateur</th><th align="left">Catégorie</th>'
                '<th align="left">Actuel</th><th align="left">Cible</th>'
                '<th align="left">Statut</th><th align="left">Tendance</th>'
                '<th align="left">Commentaire</th></tr>']
        for k in kpis:
            statut = (k or {}).get("statut", "inconnu")
            color = KPI_COLORS.get(statut, "9ca3af")
            rows.append(
                "<tr>"
                f'<td>{_esc(k.get("nom"))}</td>'
                f'<td>{_esc(k.get("categorie"))}</td>'
                f'<td>{_esc(k.get("valeur_actuelle"))}</td>'
                f'<td>{_esc(k.get("valeur_cible"))}</td>'
                f'<td><b style="color:#{color};">{_esc(KPI_LABELS.get(statut, statut))}</b></td>'
                f'<td>{_esc(k.get("tendance"))}</td>'
                f'<td>{_esc(k.get("commentaire"))}</td>'
                "</tr>"
            )
        body.append(
            '<table border="1" cellpadding="4" cellspacing="0" '
            'style="font-size:8.5pt; color:#374151; width:100%; border-color:#d1d5db;">'
            + "".join(rows) + "</table>"
        )

    elif analysis_type == "copil":
        body.append(_html_section_title("État d'avancement"))
        body.append(f'<p style="font-size:10pt; color:#374151;">{_esc(result.get("etat_avancement") or "—")}</p>')

        body.append(_html_section_title("Points clés"))
        body.append(_html_bullets(_as_list(result.get("points_cles"))))

        body.append(_html_section_title("Alertes"))
        alertes = _as_list(result.get("alertes"))
        if alertes:
            rows = ['<tr style="background-color:#f3f4f6;">'
                    '<th align="left">Type</th><th align="left">Priorité</th>'
                    '<th align="left">Description</th></tr>']
            for a in alertes:
                prio = (a or {}).get("priorite", "moyenne")
                color = RISK_COLORS.get("critique" if prio == "critique" else "eleve" if prio == "haute" else "modere", "6366f1")
                rows.append(
                    "<tr>"
                    f'<td>{_esc(ALERTE_LABELS.get(a.get("type"), a.get("type")))}</td>'
                    f'<td><b style="color:#{color};">{_esc(PRIORITE_LABELS.get(prio, prio))}</b></td>'
                    f'<td>{_esc(a.get("description"))}</td>'
                    "</tr>"
                )
            body.append(
                '<table border="1" cellpadding="4" cellspacing="0" '
                'style="font-size:9pt; color:#374151; width:100%; border-color:#d1d5db;">'
                + "".join(rows) + "</table>"
            )
        else:
            body.append('<p style="font-size:10pt; color:#9ca3af; font-style:italic;">Aucune alerte.</p>')

        body.append(_html_section_title("Décisions attendues"))
        body.append(_html_bullets(_as_list(result.get("decisions_attendues"))))

        body.append(_html_section_title("Plan d'actions"))
        actions = _as_list(result.get("plan_actions"))
        if actions:
            rows = ['<tr style="background-color:#f3f4f6;">'
                    '<th align="left">Action</th><th align="left">Responsable</th>'
                    '<th align="left">Échéance</th><th align="left">Statut</th></tr>']
            for a in actions:
                rows.append(
                    "<tr>"
                    f'<td>{_esc(a.get("action"))}</td>'
                    f'<td>{_esc(a.get("responsable"))}</td>'
                    f'<td>{_esc(a.get("echeance"))}</td>'
                    f'<td>{_esc(STATUT_ACTION_LABELS.get(a.get("statut"), a.get("statut")))}</td>'
                    "</tr>"
                )
            body.append(
                '<table border="1" cellpadding="4" cellspacing="0" '
                'style="font-size:9pt; color:#374151; width:100%; border-color:#d1d5db;">'
                + "".join(rows) + "</table>"
            )
        else:
            body.append('<p style="font-size:10pt; color:#9ca3af; font-style:italic;">Aucune action planifiée.</p>')

        body.append(_html_section_title("Résumé exécutif"))
        body.append(f'<p style="font-size:10pt; color:#374151;">{_esc(result.get("resume_executif") or "—")}</p>')

    body.append(
        f'<p style="font-size:8pt; color:#9ca3af; font-style:italic; '
        f'margin-top:20px; border-top:1px solid #e5e7eb; padding-top:6px;">{_esc(GENERATED_BY)}</p>'
    )
    return f'<html><body style="font-family:sans-serif;">{"".join(body)}</body></html>'


def _build_pdf(analysis_type: str, result: dict, project_name: str, date_str: str) -> bytes:
    html_str = _build_html(analysis_type, result, project_name, date_str)
    stream = io.BytesIO()
    story = fitz.Story(html=html_str)
    writer = fitz.DocumentWriter(stream)
    mediabox = fitz.paper_rect("a4")
    where = mediabox + (45, 45, -45, -45)
    more = 1
    while more:
        device = writer.begin_page(mediabox)
        more, _ = story.place(where)
        story.draw(device)
        writer.end_page()
    writer.close()
    return stream.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
#  GÉNÉRATION DOCX (python-docx)
# ══════════════════════════════════════════════════════════════════════════════
def _docx_header(doc, label: str, project_name: str, date_str: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(label.upper())
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(BRAND)
    run.bold = True

    title = doc.add_heading(project_name, level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    sub = doc.add_paragraph()
    r = sub.add_run(f"Document généré le {date_str}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)


def _docx_bullets(doc, items: list) -> None:
    if not items:
        doc.add_paragraph("Aucun élément.", style="Intense Quote")
        return
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def _docx_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if val is None else str(val)


def _build_docx(analysis_type: str, result: dict, project_name: str, date_str: str) -> bytes:
    doc = DocxDocument()
    label = TYPE_LABELS.get(analysis_type, analysis_type)
    _docx_header(doc, label, project_name, date_str)

    if analysis_type == "document":
        doc.add_heading("Résumé", level=1)
        doc.add_paragraph(result.get("resume") or "—")
        for key, title in [
            ("objectifs", "Objectifs"),
            ("contraintes", "Contraintes"),
            ("hypotheses", "Hypothèses"),
            ("parties_prenantes", "Parties prenantes"),
            ("points_cles", "Points clés"),
        ]:
            doc.add_heading(title, level=1)
            _docx_bullets(doc, _as_list(result.get(key)))

    elif analysis_type == "risks":
        risks = _as_list(result.get("risks"))
        doc.add_heading(f"Synthèse — {len(risks)} risque(s)", level=1)
        if result.get("resume"):
            doc.add_paragraph(result.get("resume"))
        _docx_table(
            doc,
            ["ID", "Description", "Niveau", "P", "I", "Score", "Catégorie", "Mitigation"],
            [[
                r.get("id"), r.get("description"),
                RISK_LABELS.get(r.get("niveau"), r.get("niveau")),
                r.get("probabilite"), r.get("impact"), r.get("score"),
                r.get("categorie"), r.get("mitigation"),
            ] for r in risks],
        )

    elif analysis_type == "kpi":
        kpis = _as_list(result.get("kpis"))
        score = result.get("score_global") or {}
        doc.add_heading("Score global de santé", level=1)
        p = doc.add_paragraph()
        run = p.add_run(f"{score.get('valeur', 0)}/100")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor.from_string(BRAND)
        doc.add_paragraph(score.get("interpretation") or "—")
        if result.get("resume"):
            doc.add_paragraph(result.get("resume"))
        doc.add_heading(f"Indicateurs — {len(kpis)} KPI", level=1)
        _docx_table(
            doc,
            ["Indicateur", "Catégorie", "Actuel", "Cible", "Statut", "Tendance", "Commentaire"],
            [[
                k.get("nom"), k.get("categorie"), k.get("valeur_actuelle"),
                k.get("valeur_cible"), KPI_LABELS.get(k.get("statut"), k.get("statut")),
                k.get("tendance"), k.get("commentaire"),
            ] for k in kpis],
        )

    elif analysis_type == "copil":
        doc.add_heading("État d'avancement", level=1)
        doc.add_paragraph(result.get("etat_avancement") or "—")

        doc.add_heading("Points clés", level=1)
        _docx_bullets(doc, _as_list(result.get("points_cles")))

        doc.add_heading("Alertes", level=1)
        alertes = _as_list(result.get("alertes"))
        if alertes:
            _docx_table(
                doc, ["Type", "Priorité", "Description"],
                [[
                    ALERTE_LABELS.get(a.get("type"), a.get("type")),
                    PRIORITE_LABELS.get(a.get("priorite"), a.get("priorite")),
                    a.get("description"),
                ] for a in alertes],
            )
        else:
            doc.add_paragraph("Aucune alerte.", style="Intense Quote")

        doc.add_heading("Décisions attendues", level=1)
        _docx_bullets(doc, _as_list(result.get("decisions_attendues")))

        doc.add_heading("Plan d'actions", level=1)
        actions = _as_list(result.get("plan_actions"))
        if actions:
            _docx_table(
                doc, ["Action", "Responsable", "Échéance", "Statut"],
                [[
                    a.get("action"), a.get("responsable"), a.get("echeance"),
                    STATUT_ACTION_LABELS.get(a.get("statut"), a.get("statut")),
                ] for a in actions],
            )
        else:
            doc.add_paragraph("Aucune action planifiée.", style="Intense Quote")

        doc.add_heading("Résumé exécutif", level=1)
        doc.add_paragraph(result.get("resume_executif") or "—")

    footer = doc.add_paragraph()
    fr = footer.add_run(GENERATED_BY)
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
#  GÉNÉRATION XLSX (openpyxl)
# ══════════════════════════════════════════════════════════════════════════════
_HEADER_FILL = PatternFill(start_color=BRAND, end_color=BRAND, fill_type="solid")
_HEADER_FONT = Font(color="FFFFFF", bold=True, size=11)
_TITLE_FONT = Font(color=BRAND, bold=True, size=16)
_WRAP = Alignment(wrap_text=True, vertical="top")
_THIN = Side(style="thin", color="D1D5DB")
_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)


def _xlsx_write_table(ws, start_row: int, headers: list[str], rows: list[list], widths: list[int]) -> int:
    for col, (h, w) in enumerate(zip(headers, widths), start=1):
        cell = ws.cell(row=start_row, column=col, value=h)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = Alignment(vertical="center")
        cell.border = _BORDER
        ws.column_dimensions[get_column_letter(col)].width = w
    r = start_row + 1
    for row in rows:
        for col, val in enumerate(row, start=1):
            cell = ws.cell(row=r, column=col, value="" if val is None else val)
            cell.alignment = _WRAP
            cell.border = _BORDER
        r += 1
    return r


def _xlsx_header(ws, label: str, project_name: str, date_str: str) -> int:
    ws["A1"] = project_name
    ws["A1"].font = _TITLE_FONT
    ws["A2"] = label
    ws["A2"].font = Font(color="6B7280", size=11, bold=True)
    ws["A3"] = f"Document généré le {date_str}"
    ws["A3"].font = Font(color="9CA3AF", size=9, italic=True)
    return 5  # première ligne de contenu


def _build_xlsx(analysis_type: str, result: dict, project_name: str, date_str: str) -> bytes:
    wb = Workbook()
    ws = wb.active
    label = TYPE_LABELS.get(analysis_type, analysis_type)
    ws.title = label[:31]
    row = _xlsx_header(ws, label, project_name, date_str)

    if analysis_type == "document":
        ws.cell(row=row, column=1, value="Résumé").font = Font(bold=True, color=BRAND)
        ws.cell(row=row + 1, column=1, value=result.get("resume") or "—").alignment = _WRAP
        ws.column_dimensions["A"].width = 40
        ws.column_dimensions["B"].width = 60
        r = row + 3
        for key, title in [
            ("objectifs", "Objectifs"), ("contraintes", "Contraintes"),
            ("hypotheses", "Hypothèses"), ("parties_prenantes", "Parties prenantes"),
            ("points_cles", "Points clés"),
        ]:
            r = _xlsx_write_table(ws, r, [title], [[i] for i in _as_list(result.get(key))] or [["—"]], [80])
            r += 1

    elif analysis_type == "risks":
        risks = _as_list(result.get("risks"))
        if result.get("resume"):
            ws.cell(row=row, column=1, value="Synthèse").font = Font(bold=True, color=BRAND)
            ws.cell(row=row + 1, column=1, value=result.get("resume")).alignment = _WRAP
            row += 3
        _xlsx_write_table(
            ws, row,
            ["ID", "Description", "Probabilité", "Impact", "Score", "Niveau", "Catégorie", "Mitigation"],
            [[
                r.get("id"), r.get("description"), r.get("probabilite"), r.get("impact"),
                r.get("score"), RISK_LABELS.get(r.get("niveau"), r.get("niveau")),
                r.get("categorie"), r.get("mitigation"),
            ] for r in risks],
            [10, 45, 11, 9, 8, 12, 14, 50],
        )

    elif analysis_type == "kpi":
        kpis = _as_list(result.get("kpis"))
        score = result.get("score_global") or {}
        ws.cell(row=row, column=1, value="Score global").font = Font(bold=True, color=BRAND)
        ws.cell(row=row, column=2, value=f"{score.get('valeur', 0)}/100").font = Font(bold=True)
        ws.cell(row=row + 1, column=1, value="Interprétation")
        ws.cell(row=row + 1, column=2, value=score.get("interpretation") or "—").alignment = _WRAP
        row += 3
        _xlsx_write_table(
            ws, row,
            ["ID", "Indicateur", "Catégorie", "Valeur actuelle", "Cible", "Unité", "Statut", "Tendance", "Commentaire"],
            [[
                k.get("id"), k.get("nom"), k.get("categorie"), k.get("valeur_actuelle"),
                k.get("valeur_cible"), k.get("unite"), KPI_LABELS.get(k.get("statut"), k.get("statut")),
                k.get("tendance"), k.get("commentaire"),
            ] for k in kpis],
            [10, 30, 14, 14, 12, 8, 10, 10, 45],
        )

    elif analysis_type == "copil":
        ws.cell(row=row, column=1, value="État d'avancement").font = Font(bold=True, color=BRAND)
        ws.cell(row=row + 1, column=1, value=result.get("etat_avancement") or "—").alignment = _WRAP
        ws.column_dimensions["A"].width = 30
        r = row + 3
        r = _xlsx_write_table(ws, r, ["Points clés"], [[i] for i in _as_list(result.get("points_cles"))] or [["—"]], [80])
        r += 1
        alertes = _as_list(result.get("alertes"))
        r = _xlsx_write_table(
            ws, r, ["Type", "Priorité", "Description"],
            [[
                ALERTE_LABELS.get(a.get("type"), a.get("type")),
                PRIORITE_LABELS.get(a.get("priorite"), a.get("priorite")),
                a.get("description"),
            ] for a in alertes] or [["—", "—", "—"]],
            [14, 12, 60],
        )
        r += 1
        r = _xlsx_write_table(ws, r, ["Décisions attendues"], [[i] for i in _as_list(result.get("decisions_attendues"))] or [["—"]], [80])
        r += 1
        actions = _as_list(result.get("plan_actions"))
        r = _xlsx_write_table(
            ws, r, ["Action", "Responsable", "Échéance", "Statut"],
            [[
                a.get("action"), a.get("responsable"), a.get("echeance"),
                STATUT_ACTION_LABELS.get(a.get("statut"), a.get("statut")),
            ] for a in actions] or [["—", "—", "—", "—"]],
            [50, 20, 14, 12],
        )
        r += 1
        ws.cell(row=r, column=1, value="Résumé exécutif").font = Font(bold=True, color=BRAND)
        ws.cell(row=r + 1, column=1, value=result.get("resume_executif") or "—").alignment = _WRAP

    stream = io.BytesIO()
    wb.save(stream)
    return stream.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
#  POINT D'ENTRÉE
# ══════════════════════════════════════════════════════════════════════════════
def build_export(
    analysis_type: str,
    result: dict | None,
    project_name: str,
    created_at: Any,
    fmt: str,
) -> tuple[bytes, str, str]:
    """
    Construit le fichier d'export.

    Retourne (contenu_octets, media_type, nom_de_fichier).
    Lève ValueError si le format est inconnu.
    """
    fmt = (fmt or "").lower()
    if fmt not in ALLOWED_FORMATS:
        raise ValueError(f"Format non supporté : {fmt}. Formats disponibles : {', '.join(ALLOWED_FORMATS)}")

    result = result or {}
    project_name = project_name or "Projet"
    date_str = _fmt_date(created_at)

    if fmt == "pdf":
        content = _build_pdf(analysis_type, result, project_name, date_str)
    elif fmt == "docx":
        content = _build_docx(analysis_type, result, project_name, date_str)
    else:  # xlsx
        content = _build_xlsx(analysis_type, result, project_name, date_str)

    filename = _safe_filename(project_name, analysis_type, fmt)
    return content, MEDIA_TYPES[fmt], filename
